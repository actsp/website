import streamlit as st
from streamlit_chat import message
from PIL import Image # Lib para carregar imagem no Streamlit
import matplotlib.pyplot as plt
import pandas as pd
import io
import os
from gtts import gTTS #Lib para Conversão Text2Voice. Em seguida pode usar Gemini para converter voice para texto
import google.generativeai as genai
from openai import OpenAI
from datetime import datetime
from datetime import date
import pytz
import urllib3
from urllib3 import request
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os.path
from wordcloud import WordCloud, STOPWORDS
#NLP Package
from enelvo.normaliser import Normaliser
norm = Normaliser(tokenizer='readable')
Pergunta = False
import sqlite3

import webbrowser
from ACT_CNC_lib import *


#MENU E CONFIGURAÇÕES DA PÁGINA
ajuda = "mailto:semova.app@gmail.com" 
bug = "mailto:semova.app@gmail.com"
sobre = '''
        AGENDAR Evento_Teste 26 06 2024 1800 26 06 2024 1900
        
        **Desenvolvido por Massaki de O. Igarashi e Equipe**
                                        
        '''
            
icone = "©️"
st.set_page_config(layout="wide", 
                page_title="SeMova v0.1 (By. Massaki)",
                initial_sidebar_state = "auto",
                menu_items={
                    'Get Help': (ajuda),
                    'Report a bug': (bug),
                    'About': (sobre)},
                page_icon=icone)

datetime_br= datetime.now(pytz.timezone('America/Sao_Paulo'))
t = datetime_br.strftime('%d/%m/%Y %H:%M:%S')
data_atual = datetime_br.strftime('%d/%m/%Y')
hora_atual = datetime_br.strftime('%H:%M:%S')

# Função para conectar ao banco de dados
def create_connection(db_file):
    conn = None
    try:
        conn = sqlite3.connect(db_file)
    except sqlite3.Error as e:
        st.error(f"Erro ao conectar ao banco de dados: {e}")
    return conn
    
           
#Exibe todos os registros da Base Dados User.db
def mostrar_registros():
    #Inicializa_CRUDuser:
    conn = create_connection('User.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS User (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                nome TEXT,
                keyGEMINI TEXT,
                keyOPENAI TEXT,
                situacao TEXT)''')
    c.execute('''SELECT * FROM User''')
    registros = c.fetchall()
    if registros:
        df = pd.DataFrame(registros, columns=['ID', 'NOME', 'MAIL', 'SENHA', 'CHAVE_GEMINI', 'CHAVE_OPENAI', 'SITUACAO'])
        #st.dataframe(df)
        CHAVE_GEMINI = df['CHAVE_GEMINI'][0]
        CHAVE_OPENAI = df['CHAVE_OPENAI'][0]
        return CHAVE_GEMINI, CHAVE_OPENAI
    else:
        return '0000000000000000000000000000000000000'
        st.write('Não há registros no banco de dados.')


#FUNÇÕES DA BASE DE DADOS Pesq.db
# Adiciona um novo registro à tabela
def ADD_registro(USUARIO, PESQUISA, DATA):
    #InicializarCRUDpesq:
    connection = create_connection('Pesq.db')
    sql = connection.cursor()
    sql.execute('''CREATE TABLE IF NOT EXISTS Pesq (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                usuario TEXT,
                pesquisa TEXT,
                data DATE)''')
    sql.execute('''INSERT INTO Pesq (usuario, pesquisa, data) VALUES (?, ?, ?)''',(USUARIO, PESQUISA, DATA))
    connection.commit()
    connection.close()

# Exclui um registro da tabela
def DEL_registro(id):
    connection = create_connection('Pesq.db')
    sql = connection.cursor()
    sql.execute('''DELETE FROM Pesq WHERE id = ?''', (id,))
    connection.commit()
    connection.close()

# Exibe todos os registros da tabela
def MOSTRAR_registros():
    connection = create_connection('Pesq.db')
    sql = connection.cursor()
    sql.execute('''SELECT * FROM Pesq''')
    registros = sql.fetchall()
    return registros

def mostrar_processo():
    connPROC = sqlite3.connect('Atualizacao_Processos.db')
    sqlPROC = connPROC.cursor()
    sqlPROC.execute('''SELECT * FROM Processos''')
    registrosPROC = sqlPROC.fetchall()
    return registrosPROC

def pesquisar_processo(CODIGO):
    connPROC = create_connection('Atualizacao_Processos.db')
    sqlPROC = connPROC.cursor()
    sqlPROC.execute("SELECT * FROM Processos WHERE Codigo=?", (CODIGO,))
    rows = sqlPROC.fetchall()
    return rows

def api_openai(prompt):   
    OPENAI_KEY = mostrar_registros()[1]
    client = OpenAI(api_key=OPENAI_KEY)
    completion = client.chat.completions.create(
      #model="gpt-3.5-turbo",
      model="gpt-4o",
      messages=[
        {"role": "user", "content": str(prompt)}
      ]
    )
    return completion.choices[0].message.content
    
def api_gemini(QUESTION):
    GEMINI_KEY = mostrar_registros()[0]
    genai.configure(api_key=GEMINI_KEY)

    # Set up the model
    generation_config = {
      "temperature": 0.9,
      "top_p": 1,
      "top_k": 1,
      "max_output_tokens": 2048,
    }
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(QUESTION)
    resp = response.text    
    return resp

def Nuvem_de_Palavras(msg):
    stopwords = set(STOPWORDS)
    stopwords.update(["ao", "da", "das", "de", "e", "E", "é", "É", "em", "Em", "meu", "nao", "não", "o", "ou", "os",  "para", "que", "que ", "Que", "ser", "só", "Te", "ter", "um", "você"])
    # gerar uma wordcloud
    wordcloud = WordCloud(stopwords=stopwords,
                          background_color="white",
                          width=1280, height=720).generate(msg)
    #resp = msg.replace("*", "")
    #resp = msg.replace("**", "")
    return wordcloud
def Ler_DOCx(file_path):
    DOC = Document(file_path)
    TEXTO = ''
    #TITULO = st.text_input("TÍTULO: ", DOC.paragraphs[0].text)
    Titulo = DOC.paragraphs[0].text
    n = len(DOC.paragraphs)
    for i in range(1, n):
        TEXTO+=DOC.paragraphs[i].text + '\n'
        #st.write(DOC.paragraphs[i].text)
    return Titulo, TEXTO

def main():

        
    #BARRA LATERAL
    image = Image.open('LogoNovo.jpeg')    
    st.sidebar.image(image, width=300)
    #st.image(ArqPNG, width=600, caption='Label da Figura')
    st.sidebar.divider()

    #audio_file0 = open('Generactiva.mp3', 'rb')
    #audio_bytes0 = audio_file0.read()
    #st.sidebar.audio(audio_bytes0, format='audio/ogg',start_time=0)
    video_file = open("SeMova.mp4", "rb")
    video_bytes = video_file.read()
    st.sidebar.video(video_bytes)
    st.sidebar.image("QRcode.png", caption="Sunrise by the mountains")
    Rodape = '<p style="font-weight: bolder; color:white; font-size: 16px;">Desenvolvedor: Massaki Igarashi / Integrante Equipe SeMova.</p>'
    st.sidebar.markdown(Rodape, unsafe_allow_html=True)
    mystyle0 =   '''<style> p{text-align:center;}</style>'''
    st.sidebar.markdown(mystyle0, unsafe_allow_html=True) 

    
    #============================== TÍTULO PRINCIPAL DO WEB APP ===========================================================================
    Colunas0 = st.columns(2)    
    with Colunas0[0]: 
        with st.container(height=100, border=True):
            Colunas1 = st.columns(2) 
            with Colunas1[0]:
                with st.container(height=100, border=False):
                    Titulo_Principal = '<p style="font-weight: bolder; color:#f55050; font-size: 42px;">SeMova</p>'    
                    st.markdown(Titulo_Principal, unsafe_allow_html=True)
                    st.markdown(mystyle0, unsafe_allow_html=True)    
            with Colunas1[1]:  
                #Robo = "imgs/stuser.png"
                #with st.chat_message("user", avatar=Robo):
                #    Sub_Titulo = '<p style="font-weight: bolder; color:gray; font-size: 22px;">Sua assistente!</p>'
                #    st.markdown(Sub_Titulo, unsafe_allow_html=True)        
                #    st.markdown(mystyle0, unsafe_allow_html=True) 
                with st.container(height=100, border=False):
                    Sub_Titulo = '<p style="font-weight: bolder; color: #1c5c46; font-size: 36px;">Cuide-se</p>'
                    st.markdown(Sub_Titulo, unsafe_allow_html=True)        
                    st.markdown(mystyle0, unsafe_allow_html=True) 
    with Colunas0[1]: 
        with st.container(height=100, border=True):
            st.metric("Data e Hora atual:", t)
 
    #================================================= CARDS de COTAÇÕES, NOTÍCIAS E PREVISÃO DO TEMPO ======================================
    CNC() #Faz a Chamada à Biblioteca CNC (Cotações, Notícias e Clima) para poder obter dados a partir dos seus atributos
    Noticia_Selecionada = norm.normalise(CNC.Noticia_Selecionada)
    Link_Selecionado = CNC.Link_Selecionado
    col1 = st.columns(2)
    ##st.metric("Titulo_Superior", "Valor", "Variação")
    
    
    #with col1[0]:
    #    col2 = st.columns(2)
    #    with col2[0]:
    #        with st.container(height=None, border=True):
    #            US_D = "U$ " + str(CNC.USD_DataHora)
    #            st.metric(US_D, "R$ " + str(round(float(CNC.USD_Cotacao), 2)), str(CNC.USD_Variacao)+"%")
    #    with col2[1]:
    #        with st.container(height=None, border=True):
    #            EU = "EU$ " + str(CNC.EUR_DataHora)
    #            st.metric(EU,  "R$ " + str(round(float(CNC.EUR_Cotacao), 2)), str(CNC.EUR_Variacao)+"%")
    #with col1[1]:
    #    col3 = st.columns(2)
    #    with col3[0]:
    #        with st.container(height=None, border=True):
    #            CNY = "CNY$ " + str(CNC.CNY_DataHora)
    #            st.metric(CNY,  "R$ " + str(round(float(CNC.CNY_Cotacao), 2)), str(CNC.CNY_Variacao)+"%")
    #    with col3[1]:
    #        with st.container(height=None, border=True):                
    #            if float(CNC.Umidade)<40:
    #                umidade = -1*float(CNC.Umidade)
    #            else:
    #                umidade = float(CNC.Umidade)
    #            st.metric("JDI: " +  CNC.Descricao_clima_En, str(CNC.Temp_C) + "ºC", str(CNC.Umidade) + "%" )
                
    noticia = En2Pt(Noticia_Selecionada[:120])
    with st.expander("Notícias e destaque 🌎: "):
        #st.link_button(noticia, Link_Selecionado)
        AtivaNews = True
        if AtivaNews:
            URL=str(f"[{noticia}]({Link_Selecionado})")
            st.write(URL)
        AtivaNews = st.button("↪️Atualizar Notícia!", help="Atualiza notícia", type="secondary")
        #st.write("As cinco notícias mais listadas:")
        #st.write(CNC.CincoMais)
    st.divider()
    #==============================================================================================================================

        
    tab1, tab2, tab3 = st.tabs(["Persona", "Consulta", "Resumo"]) 
    with tab1:         
        Selecao_Persona = st.radio("Persona?", ["Pedro", "Carla", "Generica"], captions=["Persona Pedro.","Persona Carla.", "Pergunta Generica.",],)
    with tab2: 
        Colunas1 = st.columns(2)
        with Colunas1[0]:
            if Selecao_Persona == "Pedro":
                Nome = st.text_input("Digite seu Nome aqui 👇", "Pedro")  
                Idade = st.text_input("Digite sua Idade aqui 👇", "28")  
                Peso = st.text_input("Digite seu Peso aqui 👇", "70") 
                Altura = st.text_input("Digite sua Altura aqui 👇", "1.65") 
                Endereco =  st.text_input("Digite seu Endereço aqui 👇", "Jundiaí - SP")
                Comportamento = st.text_input("Breve resumo de sua rotina 👇", "Vivo com meus pais e trabalho 10 horas por dia. Tenho uma vida que sobra um pouco de tempo para lazer, gosto de séries e esportes. Tenho colegas que praticam esportes que gosto, porém nem sempre estão disponíveis junto a mim")
                Necessidades = st.text_input("Principais necessidades aqui 👇", "Encontrar pessoas que gostem do mesmo tipo de atividade física que eu. Encontrar uma comunidade de prática de esportes") 
                Sexo = st.selectbox("Qual o seu Sexo?",("Masculino", "Feminino", "Não Binário"),)
                Pratica_Esportes = st.selectbox("Pratica Esportes",("Não", "Sim"),)    
                Prompt = f"Me chamo {Nome}, tenho {Idade} anos, {Peso} kg, {Altura} metros de altura, sou do sexo {Sexo}, {Pratica_Esportes} pratico esportes, preciso {Necessidades}; por isso quero que me forneça uma recomendação de esportes que devo começar a praticar baseado neste meu perfil. Ah, quero que considere também o meu comportamento a seguir: {Comportamento}. Não esqueça de me recomendar comunidades ou grupos que praticam este(s) esporte(s) em {Endereco}. Quero que seja acertivo referente às minhas habilidades e interesses para que o resultado seja compatível ao meu perfil fornecido." 
            
            elif Selecao_Persona == "Carla":
                Nome = st.text_input("Digite seu Nome aqui 👇", "Carla")  
                Idade = st.text_input("Digite sua Idade aqui 👇", "24")  
                Peso = st.text_input("Digite seu Peso aqui 👇", "65") 
                Altura = st.text_input("Digite sua Altura aqui 👇", "1.60") 
                Endereco =  st.text_input("Digite seu Endereço aqui 👇", "Jundiaí - SP")
                Comportamento = st.text_input("Breve resumo de sua rotina 👇", "Tenho uma vida muito corrida, não tenho conhecimento de esportes que gosto ou aconteça na região, não sei por onde começar. Recebi recomendação médica para iniciar o quanto antes minha prática esportiva")
                Necessidades = st.text_input("Principais necessidades aqui 👇", "Encontrar um esporte que seja alinhado a meus gostos, porém não conheço outras pessoas que possam me auxiliar na escolha") 
                Sexo = st.selectbox("Qual o seu Sexo?",("Masculino", "Feminino", "Não Binário"),index = 1)
                Pratica_Esportes = st.selectbox("Pratica Esportes",("Não", "Sim"),)    
                Prompt = f"Me chamo {Nome}, tenho {Idade} anos, {Peso} kg, {Altura} metros de altura, sou do sexo {Sexo}, {Pratica_Esportes} pratico esportes, preciso {Necessidades}; por isso quero que me forneça uma recomendação de esportes que devo começar a praticar baseado nas neste meu perfil. Ah, quero que considere também o meu comportamento a seguir: {Comportamento}. Não esqueça de me recomendar comunidades ou grupos que praticam este(s) esporte(s) em {Endereco}. Quero que seja acertivo referente às minhas habilidades e interesses para que tenha o resultado compatível ao meu perfil fornecido." 
            else:            
                question = st.text_input("Digite sua pergunta aqui 👇")
                Prompt = ""
            
            #NLP = st.checkbox("Ativar NLP!", help="Ativa o Processamento de Linguagem Natural", value=True)        
            #if NLP:
            #    question = norm.normalise(st.text_input("Digite sua pergunta aqui 👇", key="input"))
            #else:
            #    question = st.text_input("Digite sua pergunta aqui 👇", key="input") 
        
        with Colunas1[1]: 
            #    st.subheader(" ")
            #    st.subheader(" ")
                st.header(" ")
                st.write(Prompt)
                question = Prompt
                BTNquestion = st.button("Pesquisar", help="Realiza a Pesquisa nas bases de IA Generativa", type="primary", use_container_width=True)            
        with tab3:
            st.write(Prompt)
    if 'ai_answer' not in st.session_state:
        st.session_state['ai_answer'] = []

    if 'ai_question' not in st.session_state:
        st.session_state['ai_question'] = []
    
    if BTNquestion and question:
        USUARIO = "Massaki"
        ADD_registro(USUARIO, question, t)
        output1 = " "
        output2 = " "
        resumo = " "
        if "PROCESSO" in question:
            VetorQuestion = question.split()
            COD = VetorQuestion[1]
            results = pesquisar_processo(COD)
            dfPROC = pd.DataFrame(results, columns=['ID', 'CODIGO', 'CLIENTE', 'RESPONSAVEL', 'ATUALIZACAO', 'DATA', 'HORA', 'SITUACAO'])
            st.dataframe(dfPROC)
            if results:
                for row in results:
                    resumo = str(f"Atualização ({row[5]} {row[6]} ) sobre o processo Código: {row[1]}, referente ao Cliente {row[2]}: {row[4]}")
                    output1 = " "
                    output2 = " "
            #registrosPROC = mostrar_processo()
            #if registrosPROC:
            #CODIGO, CLIENTE, RESPONSAVEL, ATUALIZACAO, DATA, HORA, SITUACAO
            #    dfPROC = pd.DataFrame(registrosPROC, columns=['ID', 'CODIGO', 'CLIENTE', 'RESPONSAVEL', 'ATUALIZACAO', 'DATA', 'HORA', 'SITUACAO'])
            #    st.dataframe(dfPROC)
        elif "AGENDAR" in question:
            output1 = " "
            output2 = " "            
            evento = question.split()
            H1 = str(int(evento[5])+300)
            H2 = str(int(evento[9])+300)
            URL = str(f"https://www.google.com/calendar/render?action=TEMPLATE&text={evento[1]}&dates={evento[4]}{evento[3]}{evento[2]}T{H1}00Z%2F{evento[8]}{evento[7]}{evento[6]}T{H2}00Z")
            resumo = URL
            webbrowser.open_new_tab(URL)
        elif "Qual o seu nome?" in question:
            output1 = " "
            output2 = " "
            resumo = "Meu nome é Generactiva, sua Multi Assistente!"
        elif "Quantos anos você tem?" in question:
            output1 = " "
            output2 = " "
            resumo = "Eu nasci em 01 de junho de 2024, tenho apenas alguns dias!"       
        elif "Que dia é hoje?" in question:
            output1 = " "
            output2 = " "
            resumo = data_atual
        elif "Que horas são?" in question:
            output1 = " "
            output2 = " "
            resumo = hora_atual
        elif "Previsão do tempo" in question:
            output1 = " "
            output2 = " "
            resumo = CNC.cidade + ": " +  CNC.Descricao_clima_Pt + ", Temperatura = " + str(CNC.Temp_C) + "ºC, Umidade = " + str(CNC.Umidade) + "%"  
        elif "Cotação do dolar?" in question:
            output1 = " "
            output2 = " "
            resumo = str(CNC.USD_Cotacao) + ", variou " + str(CNC.USD_Variacao) + "%"
        elif "Cotação do euro?" in question:
            output1 = " "
            output2 = " "
            resumo = str(CNC.EUR_Cotacao) + ", variou " + str(CNC.EUR_Variacao) + "%"
        elif "Cotação do yuan?" in question:
            output1 = " "
            output2 = " "
            resumo = str(CNC.CNY_Cotacao) + ", variou " + str(CNC.CNY_Variacao)+"%"
        else:  
            output1 = api_gemini(question)
            output2 = api_openai(question)
            output = output1 + " " + output2
            resumo = api_gemini("Escreva um parágrafo que resuma o que está explicado a seguir: " + output)        
            output = output.lstrip("\n")
            output2 = output2.lstrip("\n")
        
        # Store the outputs
        st.session_state.ai_question.append(question)
        st.session_state.ai_answer.append("**Resumo:** \n" + resumo + " \n " + " \n **Resposta 1:** \n" + output1 + " \n " + " \n **Resposta 2:** \n" + output2)
        
        NomeArq = st.sidebar.text_input("Digite NOME do arquivo.DOCx e tecle ENTER: 👇", 'generactiva.docx')
        TITULO = st.sidebar.text_input("↪️ Título do arq.DOCX: ", "Programado por Massaki Igarashi")
        check_file = os.path.isfile(NomeArq)
        st.sidebar.write(check_file)        
        
        if check_file:
            #TITULO2 = st.sidebar.text_input("TÍTULO: ", Ler_DOCx(NomeArq)[0])
            memo1 = st.sidebar.text_area("Conteúdo: ", Ler_DOCx(NomeArq)[1])
            bio = io.BytesIO()
            #st.sidebar.download_button(label="⬇️ Download do arq.DOCx",
            #                                    data=bio.getvalue(),
            #                                    file_name='https://github.com/actsolucoesparapessoas/Generactiva/blob/master/generactiva.docx',
            #                                    mime="docx")
            with open(NomeArq, "rb") as doc_file:
                 DOC = doc_file.read()
            st.sidebar.download_button(label="⬇️ Download do arq.DOCx",
                                       data=DOC,
                                       file_name=NomeArq,
                                       mime="docx")
        else:
            TITULO2 = NomeArq
 
        
    message_history = st.empty()

    

    if st.session_state['ai_answer']:
        memo = str(st.session_state['ai_answer'][len(st.session_state['ai_answer'])-1])
        Passo2 = '<p style="font-weight: bolder; color:White; font-size: 16px;">Recomendações:</p>'
        st.markdown(Passo2, unsafe_allow_html=True)
        st.markdown(mystyle0, unsafe_allow_html=True)
        st.write(memo)
        memo = memo.replace("*", " ")
        memo2 = st.text_input("Resposta editável:", memo)
        
        Passo3 = '<p style="font-weight: bolder; color:White; font-size: 16px;">Recomendações + Áudio:</p>'
        st.markdown(Passo3, unsafe_allow_html=True)
        st.markdown(mystyle0, unsafe_allow_html=True)
        #NomeArq = st.sidebar.text_input("Digite nome do arquivo.DOCx e tecle ENTER: 👉", 'generactiva.docx')
        #check_file = os.path.isfile(NomeArq)
        #st.sidebar.write(check_file)
        
        if st.sidebar.button(label = '✔️ Salvar o arquivo .DOCx'):
            if check_file:
                #TITULO2 = st.sidebar.text_input("TÍTULO: ", Ler_DOCx(NomeArq)[0])
                #memo1 = st.sidebar.text_area("Conteúdo: ", Ler_DOCx(NomeArq)[1])
                document = Document()
                #document.add_heading(TITULO2, 0)            
                document.add_heading(TITULO, 0)       
                p = document.add_paragraph(memo1+memo2)
                p.bold = True
                p.italic = True
                #p.add_run('bold').bold = True
                #p.add_run(' and some ')
                #p.add_run('italic.').italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                document.save(NomeArq)
            else:            
                document = Document()
                document.add_heading(TITULO, 0)              
                p = document.add_paragraph(memo2)
                p.bold = True
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                document.save(NomeArq)

 
    if st.session_state['ai_answer']:
        my_expander = st.expander(label='Clique e abra o hitórico de Perguntas & Respostas! 👉')
        with my_expander:
            for i in range(len(st.session_state['ai_answer']) - 1, -1, -1):
                # This function displays Gemini response
                #Exemplos de avatar_styles: https://discuss.streamlit.io/t/what-are-the-other-avatar-style-icon-options/61528/4
                message("PERGUNTA: " + st.session_state['ai_question'][i], avatar_style="avataaars", is_user=True, key=str(i) + 'data_by_user')  
                perg = st.session_state['ai_question'][i] 
                
                    #avatar_style="avataaars",
                    #avatar_style="🤵", 
                    #avatar_style="initials",
                    #avatar_style="user",
                    #avatar_style="adventurer",
                
                #Pergunta para Gemini
                language = 'pt'     # Language in which you want to convert
                # Passing the text and language to the engine,here we have marked slow=False. Which tells the module that the converted audio should have a high speed
                #myobj1 = gTTS(text="PERGUNTA: " + str(perg), lang=language, slow=False)
                myobj1 = gTTS(text= perg, lang=language, slow=False)
                name1 = "perg" + str(i) + ".mp3"
                myobj1.save(name1)   #Saving the converted audio in a mp3 file
                # Playing the converted file
                audio_file = open(name1, 'rb')
                audio_bytes = audio_file.read()
                st.audio(audio_bytes, format='audio/ogg',start_time=0)
        
                # This function displays user ai_answer
                message(st.session_state["ai_answer"][i],
                    key=str(i),
                    #avatar_style="icons"
                    avatar_style="bottts"
                    )            
                resp = st.session_state["ai_answer"][i]
                resp = resp.replace("*", "")
                resp = resp.replace("**", "")
                
                #RESPOSTA Gemini
                # Passing the text and language to the engine,here we have marked slow=False. Which tells the module that the converted audio should have a high speed
                myobj2 = gTTS(text=resp, lang=language, slow=False)
                name2 = "resp" + str(i) + ".mp3"
                myobj2.save(name2)   #Saving the converted audio in a mp3 file
                # Playing the converted file
                audio_file2 = open(name2, 'rb')
                audio_bytes2 = audio_file2.read()
                st.audio(audio_bytes2, format='audio/ogg',start_time=0) 

if __name__ == '__main__':
    main()
